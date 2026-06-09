"""
题目导出服务
支持将题目导出为Word、PDF、Excel等格式
"""

import os
import json
from typing import List, Dict, Optional
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，Word导出功能不可用")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("警告: openpyxl未安装，Excel导出功能不可用")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("警告: reportlab未安装，PDF导出功能不可用")


class QuestionExporter:
    """题目导出服务"""
    
    def __init__(self, output_dir: str = "data/exports"):
        """
        初始化导出服务
        
        Args:
            output_dir: 导出文件存储目录
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export_to_json(self, questions: List[Dict], filename: Optional[str] = None) -> str:
        """
        导出为JSON格式
        
        Args:
            questions: 题目列表
            filename: 文件名（可选）
            
        Returns:
            文件路径
        """
        if filename is None:
            filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        filepath = self.output_dir / filename
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(questions, f, ensure_ascii=False, indent=2)
        
        return str(filepath)
    
    def export_to_txt(self, questions: List[Dict], filename: Optional[str] = None) -> str:
        """
        导出为文本格式
        
        Args:
            questions: 题目列表
            filename: 文件名（可选）
            
        Returns:
            文件路径
        """
        if filename is None:
            filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        filepath = self.output_dir / filename
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("=" * 60 + "\n")
            f.write("题目列表\n")
            f.write("=" * 60 + "\n\n")
            
            for i, q in enumerate(questions, 1):
                f.write(f"题目 {i}: {q.get('title', '无标题')}\n")
                f.write("-" * 60 + "\n")
                f.write(f"内容: {q.get('content', '')}\n")
                
                # 如果是选择题，显示选项
                if q.get('options'):
                    f.write("\n选项:\n")
                    for key, value in q['options'].items():
                        f.write(f"  {key}. {value}\n")
                
                # 显示答案
                answer = q.get('reference_answer', {})
                f.write(f"\n答案: {answer.get('content', '')}\n")
                
                if answer.get('explanation'):
                    f.write(f"解析: {answer.get('explanation')}\n")
                
                f.write("\n" + "=" * 60 + "\n\n")
        
        return str(filepath)
    
    def export_to_markdown(self, questions: List[Dict], filename: Optional[str] = None) -> str:
        """
        导出为Markdown格式
        
        Args:
            questions: 题目列表
            filename: 文件名（可选）
            
        Returns:
            文件路径
        """
        if filename is None:
            filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        
        filepath = self.output_dir / filename
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# 题目列表\n\n")
            
            for i, q in enumerate(questions, 1):
                f.write(f"## 题目 {i}: {q.get('title', '无标题')}\n\n")
                f.write(f"{q.get('content', '')}\n\n")
                
                # 如果是选择题，显示选项
                if q.get('options'):
                    f.write("**选项:**\n\n")
                    for key, value in q['options'].items():
                        f.write(f"- {key}. {value}\n")
                    f.write("\n")
                
                # 显示答案
                answer = q.get('reference_answer', {})
                f.write(f"**答案:** {answer.get('content', '')}\n\n")
                
                if answer.get('explanation'):
                    f.write(f"**解析:** {answer.get('explanation')}\n\n")
                
                if answer.get('key_points'):
                    f.write("**关键点:**\n")
                    for kp in answer['key_points']:
                        f.write(f"- {kp}\n")
                    f.write("\n")
                
                f.write("---\n\n")
        
        return str(filepath)
    
    def export_to_word(self, questions: List[Dict], filename: Optional[str] = None) -> str:
        """
        导出为Word格式
        
        Args:
            questions: 题目列表
            filename: 文件名（可选）
            
        Returns:
            文件路径
        """
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx未安装，无法导出Word格式")
        
        if filename is None:
            filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = self.output_dir / filename
        doc = Document()
        
        # 设置标题
        title = doc.add_heading('题目列表', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加题目
        for i, q in enumerate(questions, 1):
            # 题目标题
            doc.add_heading(f'题目 {i}: {q.get("title", "无标题")}', level=1)
            
            # 题目内容
            p = doc.add_paragraph(q.get('content', ''))
            p.style.font.size = Pt(12)
            
            # 如果是选择题，添加选项
            if q.get('options'):
                doc.add_paragraph('选项:', style='Heading 3')
                for key, value in q['options'].items():
                    doc.add_paragraph(f'{key}. {value}', style='List Bullet')
            
            # 答案
            answer = q.get('reference_answer', {})
            doc.add_paragraph(f'答案: {answer.get("content", "")}', style='Heading 3')
            
            if answer.get('explanation'):
                doc.add_paragraph(f'解析: {answer.get("explanation")}')
            
            # 添加分页符（除了最后一题）
            if i < len(questions):
                doc.add_page_break()
        
        doc.save(str(filepath))
        return str(filepath)
    
    def export_to_excel(self, questions: List[Dict], filename: Optional[str] = None) -> str:
        """
        导出为Excel格式
        
        Args:
            questions: 题目列表
            filename: 文件名（可选）
            
        Returns:
            文件路径
        """
        if not EXCEL_AVAILABLE:
            raise ValueError("openpyxl未安装，无法导出Excel格式")
        
        if filename is None:
            filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        filepath = self.output_dir / filename
        wb = Workbook()
        ws = wb.active
        ws.title = "题目列表"
        
        # 设置表头
        headers = ['序号', '标题', '内容', '类型', '难度', '科目', '答案', '解析']
        ws.append(headers)
        
        # 设置表头样式
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        
        for col in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 添加题目数据
        for i, q in enumerate(questions, 1):
            answer = q.get('reference_answer', {})
            options_text = ""
            if q.get('options'):
                options_text = "\n".join([f"{k}. {v}" for k, v in q['options'].items()])
            
            row = [
                i,
                q.get('title', ''),
                q.get('content', '') + (f"\n\n选项:\n{options_text}" if options_text else ""),
                q.get('type', ''),
                q.get('difficulty', ''),
                q.get('subject', ''),
                answer.get('content', ''),
                answer.get('explanation', '')
            ]
            ws.append(row)
        
        # 调整列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 8
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 20
        ws.column_dimensions['H'].width = 50
        
        wb.save(str(filepath))
        return str(filepath)
    
    def export_to_pdf(self, questions: List[Dict], filename: Optional[str] = None) -> str:
        """
        导出为PDF格式
        
        Args:
            questions: 题目列表
            filename: 文件名（可选）
            
        Returns:
            文件路径
        """
        if not PDF_AVAILABLE:
            raise ValueError("reportlab未安装，无法导出PDF格式")
        
        if filename is None:
            filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        filepath = self.output_dir / filename
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        
        # 创建样式
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#000000',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        story = []
        
        # 标题
        story.append(Paragraph("题目列表", title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # 添加题目
        for i, q in enumerate(questions, 1):
            # 题目标题
            story.append(Paragraph(f"题目 {i}: {q.get('title', '无标题')}", styles['Heading2']))
            story.append(Spacer(1, 0.2*inch))
            
            # 题目内容
            content = q.get('content', '').replace('\n', '<br/>')
            story.append(Paragraph(content, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # 如果是选择题，添加选项
            if q.get('options'):
                story.append(Paragraph("选项:", styles['Heading3']))
                for key, value in q['options'].items():
                    story.append(Paragraph(f"{key}. {value}", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # 答案
            answer = q.get('reference_answer', {})
            story.append(Paragraph(f"<b>答案:</b> {answer.get('content', '')}", styles['Normal']))
            
            if answer.get('explanation'):
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph(f"<b>解析:</b> {answer.get('explanation')}", styles['Normal']))
            
            # 添加分页符（除了最后一题）
            if i < len(questions):
                story.append(PageBreak())
        
        doc.build(story)
        return str(filepath)
